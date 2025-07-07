"""Маршруты для работы с документами и актами"""

import math
from flask import Blueprint, request, redirect, url_for, flash, send_file, current_app
from datetime import datetime
import os
from .auth_routes import get_current_user
from models import *
from macro_data import get_property_details
import utils
from utils import month_name_genitive
from docx import Document
import io
import re

document_bp = Blueprint('document', __name__)


@document_bp.route('/get_agreement/<int:user_id>', methods=['GET'])
def get_agreement(user_id):
    """Генерация соглашения для пользователя"""
    try:
        # Получаем данные пользователя из базы
        user = User.query.get_or_404(user_id)
        user_data = user.user_data  # Получаем связанные данные пользователя
        
        if not user_data:
            flash('Не найдены данные пользователя', 'error')
            return redirect(url_for('referal.profile'))

        # Определяем путь к шаблону
        template_path = os.path.join(current_app.root_path, 'documents', f"{os.getenv('AGREEMENT_DOC_NAME')}.docx")

        
        # Загружаем шаблон
        doc = Document(template_path)
        
        # Подготавливаем данные для замены (используем данные из UserData)
        current_date = datetime.now()
        formatted_date = f"{current_date.day} {month_name_genitive(current_date.month)} {current_date.year}"
        
        replacements = {
            #Дата в шапке (возможно не будет использоваться)
            'day': current_date.day,
            'month': month_name_genitive(current_date.month),
            'year': current_date.year,
            
            #Данные пользователя 1я страница
            'full_name': user_data.full_name or '',
            'passport_number': user_data.passport_number or '',
            'passport_giver': user_data.passport_giver or '',
            'passport_date': user_data.passport_date.strftime('%d.%m.%Y') if user_data.passport_date else '',
            'passport_address': user_data.passport_adress or '',
            
            #Подвал (паспорт адрес и имя тоже юзаются)
            'pinfl': user_data.pinfl.replace(' ','') or '',
            'trans_schet': user_data.trans_schet.replace(' ','') or '',
            'card_number': user_data.card_number.replace(' ','') or '',
            'bank': user_data.bank_name or '',
            'mfo': user_data.mfo or '',
            'phone': user_data.phone.replace(' ','') or '',
            'e_mail': user_data.e_mail or '',

        }

        required_fields = {
            'full_name': 'Полное имя пользователя',
            'passport_number': 'Паспортный номер',
            'passport_giver': 'Кем выдан паспорт',
            'passport_date': 'Дата выдачи паспорта',
            'passport_address': 'Прописка по паспорту',
            'pinfl': 'ПИНФЛ',
            'trans_schet': 'Расчетный счет',
            'card_number': 'Номер карты',
            'bank': 'Название банка',
            'mfo': 'МФО Банка',
            'phone': 'Телефон',
        }

        missing_fields = []
        for field_key, field_name in required_fields.items():
            value = replacements.get(field_key, '')
            if not value or str(value).strip() == '' or str(value).strip() == '0':
                missing_fields.append(field_name)
        
        if missing_fields:
            error_message = f"Невозможно сгенерировать акт. Отсутствуют обязательные поля: {', '.join(missing_fields)}"
            flash(error_message, 'error')
            return redirect(url_for('referal.profile'))


        # Заменяем плейсхолдеры в документе
        _replace_text_in_document(doc, replacements)

        # Сохраняем документ в память
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Генерируем имя файла
        safe_name = re.sub(r'[^\w\s-]', '', user_data.full_name or 'user').strip()
        filename = f"agreement_{safe_name}_{current_date.strftime('%Y%m%d')}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        current_app.logger.error(f"Error generating agreement: {str(e)}")
        flash('Ошибка при генерации соглашения', 'error')
        return redirect(url_for('referal.profile'))


@document_bp.route('/get_referal_act/<int:referal_id>', methods=['GET'])
def get_referal_act(referal_id):
    """Генерация акта для реферала"""
    current_user = get_current_user()
    if not current_user:
        flash('Пользователь не найден', 'error')
        return redirect(url_for('referal.profile'))
    try:
        # Получаем данные реферала и связанного пользователя из БД
        referal = Referal.query.get_or_404(referal_id)
        user = referal.user  # Получаем пользователя (реферера)
        user_data = user.user_data if user else None  # Получаем данные пользователя
        referal_data = referal.referal_data  # Получаем данные реферала
        
        if not user or not user_data:
            flash('Не найдены данные пользователя для данного реферала', 'error')
            return redirect(url_for('referal.profile'))
            
        if not referal_data:
            flash('Не найдены данные реферала', 'error')
            return redirect(url_for('referal.profile'))

        deals = MacroDeal.query.filter_by(contacts_buy_id=referal.contact_id)
        if not deals:
            flash('Deal not found', 'error')
            return redirect(url_for('referal.profile'))
    
        real_deal = MacroDeal()
        for deal in deals:
            if deal.deal_status_name == "Сделка проведена":
                real_deal = deal
                break
        appartment_info = get_property_details(real_deal.agreement_number)
        # Определяем путь к шаблону
        template_path = os.path.join(current_app.root_path, 'documents', f"{os.getenv('ACT_DOC_NAME')}.docx")

        agreement_date = appartment_info['agreement_date'].date()
        print(f"Agreement date: {agreement_date}")
        # Загружаем шаблон
        doc = Document(template_path)
        # Подготавливаем данные для замены (используем данные из UserData и ReferalData)
        current_date = datetime.now()
        print(f"Current date: {month_name_genitive(current_date.month)}")
        print(f"Current date: {month_name_genitive(int(agreement_date.month))}")
        # Форматируем дату для замены
        replacements = {
            'day': current_date.day,
            'month': month_name_genitive(current_date.month),
            'year': current_date.year,
            
            'full_name': user_data.full_name or '',
            'referal_name': referal_data.full_name or '',
            'referal_passport_number': referal_data.passport_number or '',
            'referal_passport_date': referal_data.passport_date.strftime('%d.%m.%Y') if referal_data.passport_date else '',
            'referal_passport_giver': referal_data.passport_giver or '',
            'contract_number': referal_data.contract_number or '',
            
            'contract_day': agreement_date.day,
            'contract_month': month_name_genitive(int(agreement_date.month)),
            'contract_year': agreement_date.year,
            
            'project_name': appartment_info['project_name'],
            'house_address': appartment_info['house_address'],
            'house_number': appartment_info['house_number'],
            'appartment_number': appartment_info['apartment_number'],
            'appartment_area': real_deal.deal_metr,
            'contract_price': appartment_info['agreement_price'],
            'withdrawal_amount': str(math.ceil(referal.withdrawal_amount/(1-float(os.getenv('NDS_PERCENT'))/100)) or 0),
            
            'referer_name': user_data.full_name or '',
            'passport_address': user_data.passport_adress or '',
            'pinfl': user_data.pinfl.replace(' ','') or '',
            'referer_phone': user_data.phone.replace(' ','') or '',
            'referer_email': user_data.e_mail or '',
        }

        # Проверяем наличие всех обязательных полей
        required_fields = {
            'full_name': 'Полное имя пользователя',
            'referal_name': 'Полное имя реферала',
            'referal_passport_number': 'Номер паспорта реферала',
            'referal_passport_date': 'Дата выдачи паспорта реферала',
            'referal_passport_giver': 'Орган выдачи паспорта реферала',
            'contract_number': 'Номер договора',
            'project_name': 'Название проекта',
            'house_address': 'Адрес дома',
            'house_number': 'Номер дома',
            'appartment_number': 'Номер квартиры',
            'appartment_area': 'Площадь квартиры',
            'contract_price': 'Цена договора',
            'withdrawal_amount': 'Сумма к выводу',
            'referer_name': 'Имя реферера',
            'passport_address': 'Адрес прописки реферера',
            'pinfl': 'ПИНФЛ реферера',
            'referer_phone': 'Телефон реферера',
        }
        
        missing_fields = []
        for field_key, field_name in required_fields.items():
            value = replacements.get(field_key, '')
            if not value or str(value).strip() == '' or str(value).strip() == '0':
                missing_fields.append(field_name)
        
        if missing_fields:
            error_message = f"Невозможно сгенерировать акт. Отсутствуют обязательные поля: {', '.join(missing_fields)}"
            flash(error_message, 'error')
            return redirect(url_for('referal.profile'))

        # Заменяем плейсхолдеры в документе
        _replace_text_in_document(doc, replacements)

        # Сохраняем документ в память
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Генерируем имя файла
        safe_referal_name = re.sub(r'[^\w\s-]', '', referal_data.full_name or 'referal').strip()
        filename = f"act_{safe_referal_name}_{current_date.strftime('%Y%m%d')}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        current_app.logger.error(f"Error generating referal act: {str(e)}")
        flash(f'Ошибка при генерации акта {e}', 'error')
        return redirect(url_for('referal.profile'))


def _replace_text_in_document(doc, replacements):
    """Заменяет плейсхолдеры в документе"""
    # Замена в параграфах
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', str(replacement))
    
    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(f'{{{placeholder}}}', str(replacement))
    
    # Замена в колонтитулах
    for section in doc.sections:
        # Верхний колонтитул
        if section.header:
            for paragraph in section.header.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', str(replacement))
        
        # Нижний колонтитул
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', str(replacement))
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', str(replacement))
    
    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(f'{{{placeholder}}}', str(replacement))
    
    # Замена в колонтитулах
    for section in doc.sections:
        # Верхний колонтитул
        if section.header:
            for paragraph in section.header.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', str(replacement))
        
        # Нижний колонтитул
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'{{{placeholder}}}', str(replacement))