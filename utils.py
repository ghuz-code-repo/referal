import logging
import os
import smtplib
from typing import List, Optional, Tuple
import requests
import re
from docx import Document
from datetime import datetime
import time
import requests, base64
import os
from werkzeug.security import generate_password_hash

from models import User

#from docx.shared import Pt

# Функция для отправки SMS-сообщения через API корпоративной PBX
def send_sms(phone_number, user_full_name):
    """
    Отправляет SMS-сообщение через Playmobile API.
    
    Args:
        phone_number (str): Номер телефона получателя (формат +998 XX XXX XX XX).
        user_full_name (str): Полное имя пользователя, который рекомендует.
    """
    # Clean the phone number - remove spaces and ensure it starts with "+998"
    clean_phone = phone_number.replace(" ", "").replace("+", "")
    if not clean_phone.startswith("998"):
        print(f"Invalid phone number format: {phone_number}")
        return False
        
    # Ensure user_full_name is properly decoded from Base64 if needed
    # (This is handled in routes.py, but added here for robustness)
    if isinstance(user_full_name, bytes):
        try:
            user_full_name = user_full_name.decode('utf-8')
        except UnicodeDecodeError:
            # Try to decode base64
            try:
                user_full_name = base64.b64decode(user_full_name).decode('utf-8')
            except:
                pass  # Keep as is if all decoding attempts fail
    
    # SMS text with recommendation
    sms_text = f"Вы были рекомендованы {user_full_name}. Вам предоставлена персональная скидка. Уточните удобное место и время встречи по номеру {os.getenv('GH_PHONE_NUMBER')}"
    
    # Playmobile API configuration
    api_url = os.getenv('SMS_API_URL')

    # Your Playmobile credentials
    username = os.getenv('SMS_API_USERNAME')
    password = os.getenv('SMS_API_PASSWORD')
    originator = os.getenv('SMS_API_ORIGINATOR')
    message_id = (datetime.now().strftime("%Y%m%d%H%M%S").encode('utf-8')).decode('utf-8')  # Unique message ID
    # Format payload according to Playmobile's requirements
    payload = {
        "messages": [
            {
                "recipient": clean_phone,  # Correctly formatted
                "message-id": message_id,
                "sms": {
                    "originator": originator,  # Ensure this is a string
                    "content": {
                        "text": sms_text
                    }
                }
            }
        ]
    }
    
    # # Set up authentication
    # auth = (username, password)
    
    # headers = {
    #     'Content-Type': 'application/json'
    # }
    
    try:
        userpass = username + ':' + password
        b64Val = base64.b64encode(userpass.encode('utf-8')).decode('utf-8')

        headers = {
            "Authorization": "Basic %s" % b64Val,
            "Content-Type": "application/json",
            "charset": "UTF-8"  # Explicitly set charset as mentioned in docs
        }

        # Print debug info
        #print(f"Sending to: {api_url}")
        #print(f"Headers: {headers}")
        #print(f"Payload: {payload}")

        response = requests.post(api_url, headers=headers, json=payload, timeout=30)

        print(f"Response status: {response.status_code}")
        print(f"Response text: {response.text}")

        if response.status_code == 200:
            print(f"SMS sent successfully to {phone_number}")
            return True
        else:
            print(f"Failed to send SMS: {response.status_code} - {response.text}")
            return False
        
    except requests.exceptions.RequestException as e:
        print(f"Error connecting to Playmobile: {e}")
        return False
    
def month_name_genitive(date):
    """Возвращает название месяца в родительном падеже"""
    # Create a fixed mapping instead of relying on system locales
    month_mapping = {
        1: 'Января', 2: 'Февраля', 3: 'Марта',
        4: 'Апреля', 5: 'Мая', 6: 'Июня',
        7: 'Июля', 8: 'Августа', 9: 'Сентября',
        10: 'Октября', 11: 'Ноября', 12: 'Декабря'
    }
    
    try:
        # Handle integer month number
        if isinstance(date, int) and 1 <= date <= 12:
            return month_mapping.get(date)
        # Handle datetime object
        elif hasattr(date, 'month'):
            return month_mapping.get(date.month)
        # Handle string month name (convert to corresponding number)
        else:
            # Simple mapping of Russian month names to numbers
            month_to_num = {
                'Январь': 1, 'Февраль': 2, 'Март': 3, 'Апрель': 4,
                'Май': 5, 'Июнь': 6, 'Июль': 7, 'Август': 8,
                'Сентябрь': 9, 'Октябрь': 10, 'Ноябрь': 11, 'Декабрь': 12
            }
            # Try to find the month number and then get the genitive form
            for name, num in month_to_num.items():
                if name in str(date):
                    return month_mapping.get(num)
            
            # If nothing matched, return as is
            return str(date)
    except Exception as e:
        print(f"Error in month_name_genitive: {e}")
        # Fallback - return the original value
        return str(date)
def get_document(type=None, **kwargs):
    if type is None:
        raise ValueError("Document type must be specified")
        
    template_path = os.path.join(os.path.dirname(__file__), 'documents', f"{type}.docx")

    try:
        doc = Document(template_path)
        
        # Function to replace placeholders in any paragraph
        def replace_placeholders_in_paragraph(paragraph):
            if not paragraph.runs:
                return
                
            # Combine all runs to reconstruct the entire paragraph text
            paragraph_text = ''.join([run.text for run in paragraph.runs])
            
            # Check if the paragraph contains any placeholders
            has_placeholders = False
            for key in kwargs.keys():
                placeholder = f"${{{key}}}"
                if placeholder in paragraph_text:
                    has_placeholders = True
                    paragraph_text = paragraph_text.replace(placeholder, str(kwargs[key]))
            
            # If placeholders were found, update all runs
            if has_placeholders:
                # Set the first run to the entire updated text
                paragraph.runs[0].text = paragraph_text
                
                # Clear the text from all other runs
                for run in paragraph.runs[1:]:
                    run.text = ""
        
        # Process paragraphs in the main document
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph)
        
        # Process paragraphs in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph)
        
        # Create filename for download
        output_filename = f'{type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        # Save to BytesIO
        from io import BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes, output_filename
    except Exception as e:
        print(f"Error generating document: {e}")
        return None, None

# def send_email(recipient_email, subject, body):
    """
    Отправляет email сообщение с использованием SMTP.
    """
    from email.mime.text import MIMEText
    import socket
    import time
    
    # Get email server settings from environment variables
    email_server = os.getenv('EMAIL_SERVER')
    email_port = os.getenv('EMAIL_SERVER_PORT')
    sender_email = os.getenv('SEND_FROM_EMAIL')
    sender_password = os.getenv('SEND_FROM_EMAIL_PASSWORD')
    
    # Try to get IP address from environment if available (bypass DNS)
    email_server_ip = os.getenv('EMAIL_SERVER_IP')
    
    # Configure basic logging
    if not logging.getLogger().handlers:
        logging.basicConfig(level=logging.INFO)
    
    # Log all configuration to help debug
    logging.info(f"Sending email to: {recipient_email}")
    logging.info(f"Email server: {email_server}:{email_port}")
    
    # Check if all required environment variables are set
    if not all([email_server, email_port, sender_email, sender_password]):
        error_msg = "Missing email configuration environment variables"
        logging.error(error_msg)
        return False
    
    # Prepare message
    msg = MIMEText(body)
    msg['Subject'] = subject
    msg['From'] = sender_email
    msg['To'] = recipient_email
    
    # List of servers to try (original server name + IP if provided)
    servers_to_try = [email_server]
    if email_server_ip:
        servers_to_try.insert(0, email_server_ip)  # Try IP address first if provided
    
    # Hard-coded IP fallbacks for common mail servers (as a last resort)
    mail_server_fallbacks = {
        'smtp.gmail.com': ['74.125.20.108', '74.125.20.109'],
        'mail.gh.uz': []  # Replace with actual IPs if known
    }
    
    if email_server in mail_server_fallbacks:
        servers_to_try.extend(mail_server_fallbacks[email_server])
    
    # Try to send with retries
    max_retries = 3
    retry_delay = 2  # seconds
    
    for attempt in range(max_retries):
        for server in servers_to_try:
            try:
                logging.info(f"Email sending attempt {attempt+1}/{max_retries} using server {server}")
                
                with smtplib.SMTP(server, int(email_port), timeout=10) as smtp:
                    try:
                        logging.info("Starting TLS")
                        smtp.starttls()
                        logging.info(f"Logging in as {sender_email}")
                        smtp.login(sender_email, sender_password)
                        logging.info(f"Sending email to {recipient_email}")
                        smtp.sendmail(sender_email, recipient_email, msg.as_string())
                        logging.info(f"Email sent successfully to {recipient_email}")
                        return True
                    except Exception as e:
                        logging.error(f"SMTP operation failed: {str(e)}")
                        continue
                        
            except (socket.gaierror, socket.timeout) as e:
                logging.warning(f"Network error with server {server}: {str(e)}")
                continue
            except Exception as e:
                logging.error(f"Unexpected error with server {server}: {str(e)}")
                continue
        
        # If we reach here, none of the servers worked in this attempt
        if attempt < max_retries - 1:
            logging.info(f"All servers failed. Retrying in {retry_delay} seconds...")
            time.sleep(retry_delay)
            retry_delay *= 2  # Exponential backoff
    
    logging.error("All email sending attempts failed")
    return False

def send_email(recipient_email, subject, body):
    """
    Отправляет email сообщение с использованием SMTP.
    """
    from email.mime.text import MIMEText
    msg = MIMEText(body)
    msg['Subject'] = subject
    msg['From'] = os.getenv('SEND_FROM_EMAIL') 
    msg['To'] = recipient_email

    try:
        with smtplib.SMTP(os.getenv('EMAIL_SERVER'), os.getenv('EMAIL_SERVER_PORT')) as server:
            server.starttls()
            server.login(os.getenv('SEND_FROM_EMAIL') , os.getenv('SEND_FROM_EMAIL_PASSWORD'))
            server.sendmail(os.getenv('SEND_FROM_EMAIL'), recipient_email, msg.as_string())
            logging.info("Email sent successfully")
    except Exception as e:
        logging.error(f"Failed to send email: {e}")
        time.sleep(10)
        send_email(recipient_email, subject, body)
        raise e


def clean_phone_number(phone: str) -> str:
    """Remove all non-digit characters except the leading + sign"""
    # Remove invalid characters like # and *
    phone = re.sub(r'[#\*]', '', phone)
    # Keep only digits and leading +, remove dots and other separators
    cleaned = re.sub(r'[^\d+]', '', phone)
    return cleaned

def remove_duplicated_digits(phone: str) -> str:
    """Remove duplicated parts from phone numbers"""
    digits = re.sub(r'\D', '', phone)
    
    # Check for duplicated halves (like 9989035745859989035745)
    if len(digits) > 20:
        half_len = len(digits) // 2
        first_half = digits[:half_len]
        second_half = digits[half_len:half_len*2]
        if first_half == second_half:
            return first_half
    
    # Check for repeated country codes (like 9989035745859989035745)
    # Pattern: 998XXXXXXXXX998XXXXXXXXX -> 998XXXXXXXXX
    if digits.startswith('998') and '998' in digits[3:]:
        first_998_pos = 0
        next_998_pos = digits.find('998', 3)
        if next_998_pos > 0:
            first_part = digits[first_998_pos:next_998_pos]
            second_part = digits[next_998_pos:]
            # If first part is valid Uzbek number length (12 digits), keep it
            if len(first_part) == 12:
                return first_part
    
    return digits

def extract_phone_numbers(raw_phone: str) -> List[str]:
    """Extract individual phone numbers from raw phone string"""
    if not raw_phone or raw_phone.lower() == 'anonymous' or raw_phone.strip() == '':
        return []
    
    # Remove duplicated digits first
    cleaned_raw = remove_duplicated_digits(raw_phone)
    
    # Split by common separators
    phones = re.split(r'[,;|]', cleaned_raw)
    extracted = []
    
    for phone in phones:
        phone = phone.strip()
        if not phone or phone.lower() == 'anonymous':
            continue
            
        # Clean phone but preserve structure for analysis
        cleaned = clean_phone_number(phone)
        if cleaned and len(cleaned) >= 9:  # Minimum reasonable phone length
            extracted.append(cleaned)
    
    return extracted

def format_998_number(phone: str) -> Optional[str]:
    """Format +998 numbers as: +998 {2digits} {3digits} {2digits} {2digits}"""
    # Handle dotted format like +998.909796141
    dotted_match = re.match(r'^\+(\d{3})\.(\d{2})(\d{3})(\d{2})(\d{2})$', phone)
    if dotted_match:
        country_code = dotted_match.group(1)
        formatted_number = f'+{country_code} {dotted_match.group(2)} {dotted_match.group(3)} {dotted_match.group(4)} {dotted_match.group(5)}'
        return formatted_number
   
    # Remove +998 prefix and get remaining digits
    if phone.startswith('+998'):
        digits = phone[4:]  # Remove +998
    elif phone.startswith('998'):
        digits = phone[3:]  # Remove 998
    else:
        digits = phone
    
    # Remove any remaining non-digit characters
    digits = re.sub(r'\D', '', digits)
    
    # Handle 13-digit numbers (999977352722 -> 99 997 73 52)
    if len(digits) == 12 and digits.startswith('99'):
        return f"+998 {digits[2:4]} {digits[4:7]} {digits[7:9]} {digits[9:11]}"
    
    # Handle 13-digit numbers starting with 998 (9989035745852 -> 99 890 35 74)
    if len(digits) == 13 and digits.startswith('998'):
        # Take first 12 digits (998 + 9 digits)
        valid_digits = digits[:12]
        return f"+998 {valid_digits[3:5]} {valid_digits[5:8]} {valid_digits[8:10]} {valid_digits[10:12]}"
    
    # Standard 9 digits after country code
    if len(digits) == 9:
        return f"+998 {digits[:2]} {digits[2:5]} {digits[5:7]} {digits[7:9]}"
    
    # Handle 11 digits (like 99893574585)
    if len(digits) == 11 and digits.startswith('99'):
        return f"+998 {digits[2:4]} {digits[4:7]} {digits[7:9]} {digits[9:11]}"
    
    # Handle 10 digits (like 9893574585 - missing one 9)
    if len(digits) == 10 and digits.startswith('98'):
        return f"+998 {digits[2:4]} {digits[4:7]} {digits[7:9]} {digits[9:11]}"
    
    return None

def format_7_number(phone: str) -> Optional[str]:
    """Format +7 numbers as: +7 {3digits} {3digits} {2digits} {2digits}"""
    # Handle dotted format
    dotted_match = re.match(r'^\+7\.(\d+)$', phone)
    if dotted_match:
        digits = dotted_match.group(1)
        # Handle long numbers like +7.99893123321123
        if len(digits) > 10:
            digits = digits[:10]  # Take first 10 digits
        if len(digits) == 10:
            return f"+7 {digits[:3]} {digits[3:6]} {digits[6:8]} {digits[8:10]}"
    
    # Remove +7 prefix and get remaining digits
    if phone.startswith('+7'):
        digits = phone[2:]  # Remove +7
    elif phone.startswith('7') and len(phone) > 1:
        digits = phone[1:]  # Remove leading 7
    else:
        digits = phone
    
    # Remove any remaining non-digit characters
    digits = re.sub(r'\D', '', digits)
    
    # Should have exactly 10 digits after country code
    if len(digits) == 10:
        return f"+7 {digits[:3]} {digits[3:6]} {digits[6:8]} {digits[8:10]}"
    
    # Handle longer numbers by taking first 10 digits
    if len(digits) > 10:
        digits = digits[:10]
        return f"+7 {digits[:3]} {digits[3:6]} {digits[6:8]} {digits[8:10]}"
    
    return None

def format_international_number(phone: str) -> Optional[str]:
    """Format international numbers (non +998 and +7) with basic formatting"""
    # Handle common patterns like +1.XXXXXXXXXX
    match = re.match(r'^\+(\d{1,3})\.(\d+)$', phone)
    if match:
        country_code = match.group(1)
        number = match.group(2)
        
        # Format based on country code
        if country_code == '1':  # US/Canada: +1 XXX XXX XX XX
            if len(number) >= 10:
                number = number[:10]  # Take first 10 digits
                return f"+1 {number[:3]} {number[3:6]} {number[6:8]} {number[8:10]}"
        elif len(country_code) == 2:  # Most European countries: +XX XXX XXX XXX
            if len(number) >= 8:
                if len(number) >= 9:
                    return f"+{country_code} {number[:3]} {number[3:6]} {number[6:]}"
                else:
                    return f"+{country_code} {number[:3]} {number[3:]}"
        elif len(country_code) == 3:  # Some countries: +XXX XXX XXX XXX
            if len(number) >= 7:
                if len(number) >= 9:
                    return f"+{country_code} {number[:3]} {number[3:6]} {number[6:]}"
                else:
                    return f"+{country_code} {number[:3]} {number[3:]}"
    
    # Standard international format
    if phone.startswith('+'):
        digits = re.sub(r'\D', '', phone[1:])  # Remove + and non-digits
        if 7 <= len(digits) <= 15:
            # Extract country code (1-3 digits)
            if digits[:1] == '1':  # US/Canada
                country_code = digits[:1]
                number = digits[1:]
                if len(number) >= 10:
                    number = number[:10]
                    return f"+{country_code} {number[:3]} {number[3:6]} {number[6:8]} {number[8:10]}"
            elif len(digits) >= 10:  # Assume 2-digit country code for others
                country_code = digits[:2]
                number = digits[2:]
                if len(number) >= 8:
                    return f"+{country_code} {number[:3]} {number[3:6]} {number[6:]}"
    
    return None

def process_phone_entry(line: str) -> Tuple[str, List[str]]:
    """Process a single line from the log and return ID and formatted phones"""
    # Extract ID and Raw Phone using regex
    id_match = re.search(r'ID: (\d+)', line)
    phone_match = re.search(r"Raw Phone: '([^']*)'", line)
    
    if not id_match or not phone_match:
        return "", []
    
    entry_id = id_match.group(1)
    raw_phone = phone_match.group(1)
    
    # Extract and format phone numbers
    phone_numbers = extract_phone_numbers(raw_phone)
    formatted_phones = []
    
    for phone in phone_numbers:
        formatted = format_phone_number(phone)
        if formatted:
            formatted_phones.append(formatted)
    
    return entry_id, formatted_phones

def format_phones_from_log(log_data: str) -> None:
    """Process the entire log data and print formatted results"""
    lines = log_data.strip().split('\n')
    
    for line in lines:
        if 'Raw Phone:' in line:
            entry_id, formatted_phones = process_phone_entry(line)
            
            if formatted_phones:
                phones_str = ', '.join(formatted_phones)
                print(f"ID: {entry_id}, Formatted Phones: {phones_str}")
            elif entry_id:
                print(f"ID: {entry_id}, Formatted Phones: No valid phones found")

import re
from collections import defaultdict
from typing import Dict, List, Tuple, Optional

def parse_unformatted_phones_file(filepath: str) -> List[Tuple[str, str]]:
    """
    Парсит файл с невалидными телефонами и возвращает список кортежей (ID, Raw Phone)
    """
    phones_data = []
    
    with open(filepath, 'r', encoding='utf-8') as file:
        for line in file:
            line = line.strip()
            if 'ID:' in line and 'Raw Phone:' in line:
                # Извлекаем ID и Raw Phone из строки
                id_match = re.search(r'ID:\s*(\d+)', line)
                phone_match = re.search(r"Raw Phone:\s*'([^']*)'", line)
                
                if id_match and phone_match:
                    contact_id = id_match.group(1)
                    raw_phone = phone_match.group(1)
                    phones_data.append((contact_id, raw_phone))
    
    return phones_data

def categorize_phone_validation_error(raw_phone: str) -> str:
    """
    Категоризирует причину непрохождения валидации телефона
    """
    if not raw_phone or raw_phone.strip() == '':
        return "EMPTY_PHONE"
    
    phone = raw_phone.strip()
    
    # Проверка на специальные значения
    if phone.lower() == 'anonymous':
        return "ANONYMOUS"
    
    # ВАЖНО: Множественные номера через запятую должны обрабатываться как отдельные номера
    # Эта функция должна анализировать только ОДИНОЧНЫЕ номера
    # Если есть запятая - это уже признак множественного номера
    if ',' in phone:
        return "MULTIPLE_PHONES_DETECTED"
    
    # Проверка на тестовые/фиктивные номера
    test_patterns = [
        r'^0+$',  # Только нули
        r'^1+$',  # Только единицы
        r'^9+$',  # Только девятки
        r'^123+',  # Начинается с 123
        r'(\d)\1{6,}',  # Одинаковые цифры подряд (7+ раз)
        r'^77777',  # Много семерок
        r'^99999',  # Много девяток
        r'^000000',  # Много нулей
    ]
    
    for pattern in test_patterns:
        if re.search(pattern, phone):
            return "TEST_FAKE_NUMBER"
    
    # Проверка на слишком короткие номера (меньше 7 цифр)
    digits_only = re.sub(r'\D', '', phone)
    if len(digits_only) < 7:
        return "TOO_SHORT"
    
    # Проверка на слишком длинные номера (больше 15 цифр)
    if len(digits_only) > 15:
        return "TOO_LONG"
    
    # Проверка на недопустимые символы
    if re.search(r'[#\*]', phone):
        return "INVALID_CHARACTERS"
    
    # Проверка на дублированные части номера
    if len(digits_only) > 10:
        half_len = len(digits_only) // 2
        first_half = digits_only[:half_len]
        second_half = digits_only[half_len:half_len*2]
        if first_half == second_half:
            return "DUPLICATED_DIGITS"
    
    # Проверка на нестандартные форматы
    uzbek_patterns = [
        r'^\+998\d{9}$',  # Стандартный узбекский формат
        r'^998\d{9}$',    # Без плюса
        r'^\d{9}$',       # Только номер без кода
    ]
    
    # Проверка международных форматов
    intl_patterns = [
        r'^\+\d{1,3}\.\d+',  # Международный с точками
        r'^\+\d{10,15}$',    # Стандартный международный
        r'^\d{10,15}$',      # Номер без плюса
    ]
    
    # Если номер начинается с +998 или 998, но не соответствует узбекскому формату
    if (phone.startswith('+998') or phone.startswith('998')) and not any(re.match(p, digits_only) for p in [r'998\d{9}', r'\d{9}']):
        return "INVALID_UZ_FORMAT"
    
    # Проверка на другие международные номера
    if phone.startswith('+') and not phone.startswith('+998'):
        # Проверяем корректность международного формата
        if re.match(r'^\+\d{7,15}$', phone.replace('.', '').replace('-', '').replace(' ', '')):
            return "VALID_INTERNATIONAL"
        else:
            return "INVALID_INTERNATIONAL"
    
    # Если дошли до сюда, номер имеет неопределенный формат
    return "UNKNOWN_FORMAT"

def analyze_phone_validation_errors(filepath: str) -> Dict[str, int]:
    """
    Анализирует файл с невалидными телефонами и возвращает статистику по категориям ошибок
    """
    phones_data = parse_unformatted_phones_file(filepath)
    error_categories = defaultdict(int)
    
    print(f"Всего обработано записей: {len(phones_data)}\n")
    
    # Примеры номеров для каждой категории
    examples = defaultdict(list)
    
    for contact_id, raw_phone in phones_data:
        # Если в номере есть запятая, разбиваем на отдельные номера
        if ',' in raw_phone:
            individual_phones = [p.strip() for p in raw_phone.split(',') if p.strip()]
            
            # Анализируем каждый номер отдельно
            for individual_phone in individual_phones:
                category = categorize_phone_validation_error(individual_phone)
                error_categories[category] += 1
                
                # Сохраняем примеры (максимум 3 для каждой категории)
                if len(examples[category]) < 3:
                    examples[category].append((contact_id, individual_phone))
        else:
            # Обычный одиночный номер
            category = categorize_phone_validation_error(raw_phone)
            error_categories[category] += 1
            
            # Сохраняем примеры (максимум 3 для каждой категории)
            if len(examples[category]) < 3:
                examples[category].append((contact_id, raw_phone))
    
    # Выводим статистику
    print("КАТЕГОРИИ ОШИБОК ВАЛИДАЦИИ ТЕЛЕФОНОВ:")
    print("=" * 60)
    
    # Сортируем по количеству (убывание)
    sorted_categories = sorted(error_categories.items(), key=lambda x: x[1], reverse=True)
    
    for category, count in sorted_categories:
        percentage = (count / len(phones_data)) * 100
        print(f"\n{category}: {count} номеров ({percentage:.1f}%)")
        
        # Описание категории
        descriptions = {
            "EMPTY_PHONE": "Пустые телефонные номера",
            "MULTIPLE_PHONES_DETECTED": "Множественные номера (обработаны отдельно)",
            "TOO_LONG": "Слишком длинные номера (>15 цифр)",
            "DUPLICATED_DIGITS": "Номера с дублированными частями",
            "INVALID_UZ_FORMAT": "Некорректный узбекский формат", 
            "TEST_FAKE_NUMBER": "Тестовые/фиктивные номера",
            "VALID_INTERNATIONAL": "Корректные международные номера",
            "TOO_SHORT": "Слишком короткие номера (<7 цифр)",
            "INVALID_CHARACTERS": "Недопустимые символы (#, *)",
            "INVALID_INTERNATIONAL": "Некорректные международные номера",
            "UNKNOWN_FORMAT": "Неопределенный формат",
            "ANONYMOUS": "Анонимные номера"
        }
        
        if category in descriptions:
            print(f"   Описание: {descriptions[category]}")
        
        # Примеры
        if examples[category]:
            print("   Примеры:")
            for ex_id, ex_phone in examples[category]:
                display_phone = ex_phone if ex_phone else "(пустой)"
                print(f"     ID {ex_id}: '{display_phone}'")
    
    return dict(error_categories)

def validate_phone_number(phone: str) -> Tuple[bool, str]:
    """
    Основная функция валидации телефонного номера
    Возвращает (is_valid, error_reason)
    """
    if not phone or phone.strip() == '':
        return False, "EMPTY_PHONE"
    
    category = categorize_phone_validation_error(phone)
    
    # Валидными считаем только определенные категории
    valid_categories = ["VALID_INTERNATIONAL"]
    
    # Дополнительная проверка для узбекских номеров
    if category == "UNKNOWN_FORMAT":
        digits = re.sub(r'\D', '', phone)
        if re.match(r'^998\d{9}$', digits) or re.match(r'^\d{9}$', digits):
            return True, "VALID_UZ_FORMAT"
    
    is_valid = category in valid_categories
    return is_valid, category

def format_phone_number(phone: str) -> Optional[str]:
    """Format a single phone number according to the specified patterns"""
    if not phone:
        return None
    
    # Clean the phone number first
    clean_phone = clean_phone_number(phone)
    
    # Handle +998 numbers (including those starting with just 998)
    if clean_phone.startswith('+998') or clean_phone.startswith('998'):
        return format_998_number(clean_phone)
    
    # Handle +7 numbers (including those starting with just 7)
    elif clean_phone.startswith('+7') or (clean_phone.startswith('7') and len(clean_phone) >= 11):
        return format_7_number(clean_phone)
    
    # Handle other international numbers
    elif clean_phone.startswith('+'):
        return format_international_number(clean_phone)
    
    # If no country code but looks like Uzbek number (11+ digits starting with 998)
    elif len(clean_phone) >= 11 and clean_phone.startswith('998'):
        return format_998_number(clean_phone)
    
    # If no country code but looks like Uzbek number (9+ digits, common mobile prefixes)
    elif len(clean_phone) >= 9 and clean_phone.isdigit():
        # Check if it starts with common Uzbek mobile prefixes
        uzbek_prefixes = ['90', '91', '93', '94', '95', '97', '98', '99', '77', '88', '33', '50', '55', '71']
        if clean_phone[:2] in uzbek_prefixes:
            # If it's exactly 9 digits, add +998
            if len(clean_phone) == 9:
                return format_998_number('+998' + clean_phone)
            # If it's 11 digits starting with 99, it might be missing 8
            elif len(clean_phone) == 11 and clean_phone.startswith('99'):
                return format_998_number(clean_phone)
            # If it's 12 digits and starts with 999, treat as 998+9digits
            elif len(clean_phone) == 12 and clean_phone.startswith('999'):
                return format_998_number(clean_phone)
    
    # If no country code but looks like Russian number (10 digits starting with 9)
    elif len(clean_phone) == 10 and clean_phone.startswith('9'):
        return format_7_number('+7' + clean_phone)
    
    # If it's 11 digits starting with 7 (Russian format without +)
    elif len(clean_phone) == 11 and clean_phone.startswith('7'):
        return format_7_number(clean_phone)
    
    return None

# Запуск анализа
if __name__ == "__main__":
    filepath = r"c:\Users\d.tolkunov\CodeRepository\AnalyticsRepo\unformatted_phones.txt"
    results = analyze_phone_validation_errors(filepath)
    
    print(f"\n\nИТОГО: {sum(results.values())} номеров распределены по {len(results)} категориям")